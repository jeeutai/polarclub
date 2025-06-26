import streamlit as st
import pandas as pd
from datetime import datetime, date, timedelta
import io
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from error_handler import error_handler

class ReportGenerator:
    def __init__(self):
        pass

    def show_report_interface(self, user):
        st.markdown("### 📊 보고서")
        if user['role'] == '선생님':
            tabs = st.tabs(["📄 활동 보고서 작성", "📋 기존 보고서", "📊 통계 보고서"])
        else:
            tabs = st.tabs(["📄 활동 보고서 작성", "📋 기존 보고서"])

        with tabs[0]:
            self.show_custom_report_form(user)
        with tabs[1]:
            self.show_existing_reports(user)
        if user['role'] == '선생님':
            with tabs[2]:
                self.show_statistics_reports(user)

    def show_custom_report_form(self, user):
        st.markdown("#### 📝 활동 보고서 작성")

        with st.form("custom_report_form"):
            col1, col2 = st.columns(2)
            with col1:
                club_name = st.text_input("🏷️ 동아리명", value="프로그래밍 동아리")
                activity_date = st.date_input("📅 활동 날짜", value=date.today())
                activity_time = st.text_input("⏰ 활동 시간", value="14:00 - 16:00")
                location = st.text_input("📍 활동 장소", value="컴퓨터실")

            with col2:
                instructor = st.text_input("👨‍🏫 지도교사", value="김선생님")
                total_members = st.number_input("👥 총 인원", min_value=1, value=15)
                present_members = st.number_input("✅ 참석 인원", min_value=1, value=12)
                weather = st.selectbox("🌤️ 날씨", ["맑음", "흐림", "비", "눈"])

            st.markdown("---")

            activity_title = st.text_input("📋 활동 제목", value="웹사이트 개발 프로젝트")
            activity_objective = st.text_area("🎯 활동 목표", 
                value="HTML, CSS, JavaScript를 활용한 반응형 웹사이트 제작 능력 향상")

            activity_content = st.text_area("📝 활동 내용", height=150,
                value="""1. 웹 개발 기초 이론 학습 (30분)
- HTML5 시맨틱 태그 구조 이해
- CSS3 플렉스박스와 그리드 레이아웃
- JavaScript DOM 조작 기법

2. 실습 활동 (60분)
- 개인 포트폴리오 웹사이트 제작
- 반응형 디자인 적용
- 상호작용 요소 구현

3. 결과 발표 및 피드백 (30분)
- 각자 제작한 웹사이트 시연
- 동료 간 코드 리뷰 진행""")

            materials = st.text_area("🛠️ 준비물 및 자료", 
                value="노트북, VS Code, Chrome 브라우저, 참고 자료집, USB")

            st.markdown("#### 👥 개별 역할 및 참여 현황")

            roles_data = []
            num_roles = st.number_input("역할 개수", min_value=1, max_value=10, value=4)

            for i in range(num_roles):
                col1, col2, col3 = st.columns([2, 2, 3])
                with col1:
                    name = st.text_input(f"이름 {i+1}", value=f"학생{i+1}", key=f"name_{i}")
                with col2:
                    role = st.text_input(f"역할 {i+1}", value=["팀장", "개발자", "디자이너", "테스터"][i] if i < 4 else "팀원", key=f"role_{i}")
                with col3:
                    performance = st.text_input(f"성과/기여도 {i+1}", 
                        value=["프로젝트 전체 관리 및 일정 조율", "메인 기능 개발 및 구현", "UI/UX 디자인 및 스타일링", "버그 테스트 및 품질 관리"][i] if i < 4 else "팀 활동 참여", 
                        key=f"perf_{i}")

                roles_data.append({"이름": name, "역할": role, "성과": performance})

            st.markdown("#### 📈 활동 결과 및 평가")

            col1, col2 = st.columns(2)
            with col1:
                achievement = st.text_area("🏆 주요 성과", 
                    value="모든 팀원이 개인 포트폴리오 웹사이트를 완성하여 실무 능력 향상")
                difficulty = st.text_area("⚠️ 어려웠던 점", 
                    value="반응형 디자인 구현 시 미디어 쿼리 적용에 어려움")

            with col2:
                improvement = st.text_area("💡 개선사항", 
                    value="다음 활동에서는 더 체계적인 단계별 가이드 제공 필요")
                next_plan = st.text_area("📋 차회 계획", 
                    value="데이터베이스 연동 및 백엔드 개발 기초 학습")

            overall_rating = st.slider("📊 전체 만족도 (1-10점)", 1, 10, 8)

            additional_notes = st.text_area("📌 특이사항 및 비고", 
                value="3명의 학생이 우수한 결과물을 완성하여 전시 예정")

            col1, col2 = st.columns(2)
            with col1:
                generate_docx = st.form_submit_button("📄 DOCX 보고서 생성", use_container_width=True)
            with col2:
                generate_pdf = st.form_submit_button("📄 PDF 보고서 생성", use_container_width=True)

        report_data = {
            "club_name": club_name, "activity_date": activity_date, "activity_time": activity_time,
            "location": location, "instructor": instructor, "total_members": total_members,
            "present_members": present_members, "weather": weather, "activity_title": activity_title,
            "activity_objective": activity_objective, "activity_content": activity_content,
            "materials": materials, "roles_data": roles_data, "achievement": achievement,
            "difficulty": difficulty, "improvement": improvement, "next_plan": next_plan,
            "overall_rating": overall_rating, "additional_notes": additional_notes
        }

        if generate_docx:
            docx_file = self.create_professional_docx_report(report_data)
            if docx_file:
                st.success("DOCX 보고서가 생성되었습니다!")
                st.download_button(
                    label="💾 DOCX 다운로드",
                    data=docx_file,
                    file_name=f"활동보고서_{club_name}_{activity_date}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        if generate_pdf:
            pdf_file = self.create_professional_pdf_report(report_data)
            if pdf_file:
                st.success("PDF 보고서가 생성되었습니다!")
                st.download_button(
                    label="💾 PDF 다운로드",
                    data=pdf_file,
                    file_name=f"활동보고서_{club_name}_{activity_date}.pdf",
                    mime="application/pdf"
                )

    def create_professional_docx_report(self, data):
        try:
            doc = Document()

            title = doc.add_heading('동아리 활동 보고서', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_heading('1. 기본 정보', level=1)
            info_table = doc.add_table(rows=4, cols=4)
            info_table.style = 'Table Grid'

            cells = info_table.rows[0].cells
            cells[0].text = '동아리명'; cells[1].text = data['club_name']
            cells[2].text = '활동일자'; cells[3].text = str(data['activity_date'])

            cells = info_table.rows[1].cells
            cells[0].text = '활동시간'; cells[1].text = data['activity_time']
            cells[2].text = '활동장소'; cells[3].text = data['location']

            cells = info_table.rows[2].cells
            cells[0].text = '지도교사'; cells[1].text = data['instructor']
            cells[2].text = '날씨'; cells[3].text = data['weather']

            cells = info_table.rows[3].cells
            cells[0].text = '총 인원'; cells[1].text = str(data['total_members'])
            cells[2].text = '참석인원'; cells[3].text = str(data['present_members'])

            doc.add_heading('2. 활동 개요', level=1)
            doc.add_paragraph(f"활동명: {data['activity_title']}", style='List Bullet')
            doc.add_paragraph(f"목표: {data['activity_objective']}", style='List Bullet')

            doc.add_heading('3. 활동 내용', level=1)
            doc.add_paragraph(data['activity_content'])

            doc.add_heading('4. 준비물 및 자료', level=1)
            doc.add_paragraph(data['materials'])

            doc.add_heading('5. 개별 역할 및 성과', level=1)
            role_table = doc.add_table(rows=1, cols=3)
            role_table.style = 'Table Grid'
            hdr_cells = role_table.rows[0].cells
            hdr_cells[0].text = '이름'; hdr_cells[1].text = '역할'; hdr_cells[2].text = '성과/기여도'

            for role in data['roles_data']:
                row_cells = role_table.add_row().cells
                row_cells[0].text = role['이름']
                row_cells[1].text = role['역할']
                row_cells[2].text = role['성과']

            doc.add_heading('6. 활동 결과 및 평가', level=1)
            doc.add_paragraph(f"주요 성과: {data['achievement']}", style='List Bullet')
            doc.add_paragraph(f"어려웠던 점: {data['difficulty']}", style='List Bullet')
            doc.add_paragraph(f"개선사항: {data['improvement']}", style='List Bullet')
            doc.add_paragraph(f"차회 계획: {data['next_plan']}", style='List Bullet')
            doc.add_paragraph(f"전체 만족도: {data['overall_rating']}/10점", style='List Bullet')

            doc.add_heading('7. 특이사항', level=1)
            doc.add_paragraph(data['additional_notes'])

            doc.add_paragraph(f"\n작성일: {datetime.now().strftime('%Y년 %m월 %d일')}")
            doc.add_paragraph(f"작성자: {data['instructor']}")

            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            return doc_io.getvalue()
        except Exception as e:
            st.error(f"DOCX 생성 오류: {e}")
            return None

    def create_professional_pdf_report(self, data):
        try:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            styles = getSampleStyleSheet()

            title_style = ParagraphStyle(
                'CustomTitle', parent=styles['Heading1'],
                fontSize=18, spaceAfter=30, alignment=1
            )

            story = []
            story.append(Paragraph("동아리 활동 보고서", title_style))
            story.append(Spacer(1, 20))

            story.append(Paragraph("1. 기본 정보", styles['Heading2']))

            basic_info = [
                ['동아리명', data['club_name'], '활동일자', str(data['activity_date'])],
                ['활동시간', data['activity_time'], '활동장소', data['location']],
                ['지도교사', data['instructor'], '날씨', data['weather']],
                ['총 인원', str(data['total_members']), '참석인원', str(data['present_members'])]
            ]

            basic_table = Table(basic_info)
            basic_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))

            story.append(basic_table)
            story.append(Spacer(1, 20))

            story.append(Paragraph("2. 활동 개요", styles['Heading2']))
            story.append(Paragraph(f"<b>활동명:</b> {data['activity_title']}", styles['Normal']))
            story.append(Paragraph(f"<b>목표:</b> {data['activity_objective']}", styles['Normal']))
            story.append(Spacer(1, 12))

            story.append(Paragraph("3. 활동 내용", styles['Heading2']))
            story.append(Paragraph(data['activity_content'].replace('\n', '<br/>'), styles['Normal']))
            story.append(Spacer(1, 12))

            story.append(Paragraph("4. 개별 역할 및 성과", styles['Heading2']))
            role_data = [['이름', '역할', '성과/기여도']]
            for role in data['roles_data']:
                role_data.append([role['이름'], role['역할'], role['성과']])

            role_table = Table(role_data, colWidths=[1.5*inch, 1.5*inch, 3*inch])
            role_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))

            story.append(role_table)
            story.append(Spacer(1, 20))

            story.append(Paragraph("5. 활동 결과 및 평가", styles['Heading2']))
            story.append(Paragraph(f"<b>주요 성과:</b> {data['achievement']}", styles['Normal']))
            story.append(Paragraph(f"<b>어려웠던 점:</b> {data['difficulty']}", styles['Normal']))
            story.append(Paragraph(f"<b>개선사항:</b> {data['improvement']}", styles['Normal']))
            story.append(Paragraph(f"<b>차회 계획:</b> {data['next_plan']}", styles['Normal']))
            story.append(Paragraph(f"<b>전체 만족도:</b> {data['overall_rating']}/10점", styles['Normal']))
            story.append(Spacer(1, 12))

            story.append(Paragraph("6. 특이사항", styles['Heading2']))
            story.append(Paragraph(data['additional_notes'], styles['Normal']))
            story.append(Spacer(1, 20))

            story.append(Paragraph(f"작성일: {datetime.now().strftime('%Y년 %m월 %d일')}", styles['Normal']))
            story.append(Paragraph(f"작성자: {data['instructor']}", styles['Normal']))

            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
        except Exception as e:
            st.error(f"PDF 생성 오류: {e}")
            return None

    def show_existing_reports(self, user):
        st.markdown("#### 📋 기존 보고서")
        st.info("저장된 보고서가 없습니다. 새로운 보고서를 작성해보세요!")

    def show_statistics_reports(self, user):
        st.markdown("#### 📊 통계 보고서")

        col1, col2, col3 = st.columns(3)
        with col1:
            error_handler.wrap_streamlit_component(st.metric, "총 보고서 수", "0")
        with col2:
            error_handler.wrap_streamlit_component(st.metric, "이번 달 활동", "0")
        with col3:
            error_handler.wrap_streamlit_component(st.metric, "평균 만족도", "0.0")